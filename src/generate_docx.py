from pathlib import Path
from docx import Document


INPUT_FILE = "data/sample.txt"
OUTPUT_FILE = "data/sample.docx"


def generate_docx_from_txt():
    input_file = Path(INPUT_FILE)
    output_file = Path(OUTPUT_FILE)

    text = input_file.read_text(encoding="utf-8")

    document = Document()
    document.add_heading("Synthetic Document Dataset", level=1)

    for block in text.split("=" * 80):
        clean_block = block.strip()

        if clean_block:
            document.add_paragraph(clean_block)

    document.save(output_file)

    print(f"DOCX generated successfully: {output_file}")


if __name__ == "__main__":
    generate_docx_from_txt()